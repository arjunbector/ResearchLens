# Import the required modules
from pdf2docx import Converter

def pdf_to_docx(pdf_file="backend/RPaper2.pdf", docx_file="research_paper.docx"):
    cv = Converter(pdf_file)
    cv.convert(docx_file)
    cv.close()
    
from docx import Document
def change_to_one_column(file_path='backend/research_paper.docx'):
    doc = Document(file_path)
    new_doc = Document()
    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size
    new_doc.save('research.docx')

file_path = 'research_paper.docx'



from docx2pdf import convert

import os
def delete_file(file_path):
    # Check if file exists
    if os.path.isfile(file_path):
        os.remove(file_path)  # Delete the file
        print(f"File {file_path} has been deleted.")
    else:
        print(f"File {file_path} does not exist.")
    
def dell():
    delete_file('research.docx')
    delete_file('research_paper.docx')
    delete_file('3542.pdf')





