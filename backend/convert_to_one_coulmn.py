from docx import Document

def change_to_one_column(file_path):
    doc = Document(file_path)
    new_doc = Document()

    # Copy the content of the old document to the new document
    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size

    new_doc.save('research.docx')

file_path = 'research_paper.docx'
change_to_one_column(file_path)